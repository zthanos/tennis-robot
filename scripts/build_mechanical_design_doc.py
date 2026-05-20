from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
IMG_DIR = DOCS / "images"
OUT = DOCS / "tennis-robot-mechanical-design.docx"


PALETTE = {
    "ink": (31, 41, 55),
    "muted": (91, 105, 124),
    "line": (126, 143, 163),
    "panel": (242, 246, 249),
    "accent": (20, 118, 130),
    "accent2": (232, 170, 66),
    "danger": (188, 76, 76),
    "green": (96, 143, 78),
    "white": (255, 255, 255),
}


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded(draw, box, fill, outline=None, width=3, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill=None, width=4):
    fill = fill or PALETTE["accent"]
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    import math

    angle = math.atan2(ey - sy, ex - sx)
    size = 14
    pts = [
        (ex, ey),
        (ex - size * math.cos(angle - 0.45), ey - size * math.sin(angle - 0.45)),
        (ex - size * math.cos(angle + 0.45), ey - size * math.sin(angle + 0.45)),
    ]
    draw.polygon(pts, fill=fill)


def label(draw, xy, text, size=24, fill=None, bold=False, anchor="mm"):
    draw.text(xy, text, fill=fill or PALETTE["ink"], font=font(size, bold), anchor=anchor)


def make_base_plan():
    path = IMG_DIR / "design-base-plan.png"
    im = Image.new("RGB", (1600, 950), PALETTE["white"])
    d = ImageDraw.Draw(im)
    label(d, (80, 60), "Βάση - κάτοψη", 34, bold=True, anchor="la")
    label(d, (80, 105), "2 drive wheels, πτυσσόμενο χερούλι, σταθεροποίηση στην εκτόξευση", 24, PALETTE["muted"], anchor="la")

    body = (370, 210, 1230, 770)
    rounded(d, body, PALETTE["panel"], PALETTE["line"], 4, 42)
    rounded(d, (470, 310, 1130, 660), (255, 255, 255), PALETTE["line"], 3, 28)
    label(d, (800, 365), "Κρυφό εσωτερικό σασί", 28, bold=True)
    label(d, (800, 410), "αλουμίνιο 2020/2040 ή plate 3-5 mm", 22, PALETTE["muted"])

    # wheels
    for x in (330, 1270):
        d.ellipse((x - 70, 420, x + 70, 620), fill=(49, 57, 67), outline=PALETTE["ink"], width=4)
        d.ellipse((x - 38, 460, x + 38, 580), fill=(130, 145, 160), outline=PALETTE["ink"], width=3)
    label(d, (330, 665), "τροχός 150-200 mm", 21)
    label(d, (1270, 665), "τροχός 150-200 mm", 21)

    # handle
    d.line((540, 180, 540, 80, 1060, 80, 1060, 180), fill=PALETTE["accent"], width=12)
    rounded(d, (645, 45, 955, 105), (232, 250, 252), PALETTE["accent"], 4, 18)
    label(d, (800, 75), "τηλεσκοπικό χερούλι", 24, PALETTE["accent"], bold=True)

    # stabilizer feet front
    for x in (540, 1060):
        rounded(d, (x - 70, 785, x + 70, 855), (255, 246, 229), PALETTE["accent2"], 4, 16)
    label(d, (800, 840), "πτυσσόμενα stabilizer feet / rubber pads", 23, PALETTE["accent2"], bold=True)

    # component zones
    rounded(d, (520, 450, 730, 600), (237, 246, 237), PALETTE["green"], 3, 18)
    label(d, (625, 505), "μπαταρία", 24, bold=True)
    label(d, (625, 540), "χαμηλά", 21, PALETTE["muted"])
    rounded(d, (770, 450, 1080, 600), (238, 246, 249), PALETTE["accent"], 3, 18)
    label(d, (925, 505), "drivers / controller", 24, bold=True)
    label(d, (925, 540), "service hatch στο πλάι", 21, PALETTE["muted"])

    label(d, (800, 905), "Κατάσταση λειτουργίας: ακουμπά σε τροχούς + stabilizers. Κατάσταση μεταφοράς: σηκώνεται από το χερούλι.", 22, PALETTE["muted"])
    im.save(path)
    return path


def make_side_section():
    path = IMG_DIR / "design-side-section.png"
    im = Image.new("RGB", (1600, 950), PALETTE["white"])
    d = ImageDraw.Draw(im)
    label(d, (80, 60), "Πλάγια τομή - διάταξη modules", 34, bold=True, anchor="la")

    # ground and body
    d.line((80, 820, 1520, 820), fill=(180, 190, 200), width=5)
    rounded(d, (260, 210, 1320, 770), PALETTE["panel"], PALETTE["line"], 4, 42)
    d.ellipse((295, 675, 455, 835), fill=(51, 58, 68), outline=PALETTE["ink"], width=4)
    d.ellipse((1145, 675, 1305, 835), fill=(51, 58, 68), outline=PALETTE["ink"], width=4)
    rounded(d, (730, 770, 920, 840), (255, 246, 229), PALETTE["accent2"], 4, 16)

    # modules
    rounded(d, (365, 585, 575, 720), (255, 255, 255), PALETTE["green"], 3, 18)
    label(d, (470, 630), "intake", 25, bold=True)
    label(d, (470, 665), "μπροστινό", 21, PALETTE["muted"])
    rounded(d, (600, 305, 850, 610), (255, 255, 255), PALETTE["accent"], 3, 22)
    label(d, (725, 345), "hopper", 25, bold=True)
    label(d, (725, 380), "40-80 μπάλες", 21, PALETTE["muted"])
    rounded(d, (900, 365, 1175, 590), (255, 255, 255), PALETTE["danger"], 3, 22)
    label(d, (1038, 410), "launcher", 25, bold=True)
    label(d, (1038, 445), "διπλοί τροχοί", 21, PALETTE["muted"])
    rounded(d, (610, 640, 1040, 735), (237, 246, 237), PALETTE["green"], 3, 18)
    label(d, (825, 688), "μπαταρία χαμηλά για κέντρο βάρους", 22, PALETTE["ink"], bold=True)

    # arrows
    arrow(d, (500, 575), (650, 510), PALETTE["green"])
    arrow(d, (820, 455), (905, 455), PALETTE["accent"])
    arrow(d, (1175, 455), (1365, 405), PALETTE["danger"], 5)
    label(d, (1395, 392), "έξοδος μπάλας", 22, PALETTE["danger"], anchor="la")

    # handle
    d.line((1220, 210, 1370, 80, 1490, 80), fill=PALETTE["accent"], width=10)
    label(d, (1400, 120), "πτυσσόμενο χερούλι", 22, PALETTE["accent"], bold=True)
    im.save(path)
    return path


def make_pickup_flow():
    path = IMG_DIR / "design-pickup-flow.png"
    im = Image.new("RGB", (1600, 900), PALETTE["white"])
    d = ImageDraw.Draw(im)
    label(d, (80, 60), "Μηχανισμός μαζέματος - ροή μπάλας", 34, bold=True, anchor="la")
    label(d, (80, 105), "Χαμηλό στόμιο, brush roller, ράμπα και elevator προς hopper", 24, PALETTE["muted"], anchor="la")

    d.line((100, 725, 1500, 725), fill=(185, 195, 205), width=5)
    # ball positions
    for x in (170, 360, 560):
        d.ellipse((x - 42, 640, x + 42, 724), fill=(210, 235, 80), outline=(130, 150, 30), width=4)
    # funnel guides
    d.polygon([(270, 640), (540, 560), (540, 625), (340, 700)], fill=(237, 246, 237), outline=PALETTE["green"])
    d.polygon([(270, 810), (540, 745), (540, 680), (340, 730)], fill=(237, 246, 237), outline=PALETTE["green"])
    label(d, (410, 530), "πλαϊνοί οδηγοί / funnel", 23, PALETTE["green"], bold=True)
    # roller
    d.ellipse((600, 560, 760, 720), fill=(70, 82, 96), outline=PALETTE["ink"], width=4)
    d.arc((615, 575, 745, 705), 35, 330, fill=PALETTE["accent2"], width=8)
    label(d, (680, 540), "brush roller", 24, bold=True)
    # ramp/elevator
    d.polygon([(760, 700), (1130, 430), (1180, 500), (820, 735)], fill=(242, 246, 249), outline=PALETTE["line"])
    for i in range(5):
        x = 810 + i * 70
        y = 690 - i * 52
        d.line((x, y, x + 55, y - 38), fill=PALETTE["accent"], width=5)
    label(d, (1035, 405), "elevator / belt με paddles", 23, PALETTE["accent"], bold=True)
    # hopper
    rounded(d, (1180, 260, 1460, 520), (255, 255, 255), PALETTE["accent"], 4, 24)
    label(d, (1320, 305), "hopper", 26, bold=True)
    for x, y in [(1250, 395), (1325, 390), (1390, 410), (1290, 465), (1360, 465)]:
        d.ellipse((x - 35, y - 35, x + 35, y + 35), fill=(210, 235, 80), outline=(130, 150, 30), width=3)
    arrow(d, (205, 595), (590, 630), PALETTE["green"])
    arrow(d, (760, 610), (1140, 420), PALETTE["accent"])
    im.save(path)
    return path


def make_intake_front_view():
    path = IMG_DIR / "design-intake-front-view.png"
    im = Image.new("RGB", (1600, 900), PALETTE["white"])
    d = ImageDraw.Draw(im)
    label(d, (80, 60), "Μπροστινή όψη intake", 34, bold=True, anchor="la")
    label(d, (80, 105), "Οδηγοί funnel, στόμιο και brush roller όπως φαίνονται από μπροστά", 24, PALETTE["muted"], anchor="la")

    d.line((120, 760, 1480, 760), fill=(185, 195, 205), width=5)
    rounded(d, (270, 160, 1330, 750), PALETTE["panel"], PALETTE["line"], 4, 42)
    rounded(d, (360, 250, 1240, 720), (255, 255, 255), PALETTE["line"], 3, 28)

    for x in (250, 1350):
        d.ellipse((x - 85, 560, x + 85, 780), fill=(49, 57, 67), outline=PALETTE["ink"], width=4)
        d.ellipse((x - 46, 610, x + 46, 735), fill=(130, 145, 160), outline=PALETTE["ink"], width=3)

    d.polygon([(405, 360), (620, 500), (620, 600), (405, 690)], fill=(237, 246, 237), outline=PALETTE["green"])
    d.polygon([(1195, 360), (980, 500), (980, 600), (1195, 690)], fill=(237, 246, 237), outline=PALETTE["green"])
    d.line((405, 360, 620, 500, 980, 500, 1195, 360), fill=PALETTE["green"], width=5)
    d.line((405, 690, 620, 600, 980, 600, 1195, 690), fill=PALETTE["green"], width=5)
    label(d, (505, 330), "αριστερός οδηγός", 22, PALETTE["green"], bold=True)
    label(d, (1095, 330), "δεξιός οδηγός", 22, PALETTE["green"], bold=True)
    label(d, (800, 465), "φαρδύ στόμιο intake", 24, PALETTE["green"], bold=True)

    rounded(d, (585, 610, 1015, 710), (245, 248, 250), PALETTE["line"], 3, 18)
    d.rounded_rectangle((620, 590, 980, 680), radius=45, fill=(70, 82, 96), outline=PALETTE["ink"], width=4)
    for x in range(650, 970, 34):
        d.line((x, 595, x - 18, 675), fill=PALETTE["accent2"], width=5)
    label(d, (800, 575), "brush roller", 24, PALETTE["accent2"], bold=True)

    for x in (690, 800, 910):
        d.ellipse((x - 36, 680, x + 36, 752), fill=(210, 235, 80), outline=(130, 150, 30), width=3)
    label(d, (800, 795), "ο roller τραβά τη μπάλα μέσα από το κεντραρισμένο throat", 22, PALETTE["muted"])

    arrow(d, (410, 250), (1190, 250), PALETTE["accent"])
    arrow(d, (1190, 250), (410, 250), PALETTE["accent"])
    label(d, (800, 225), "πλάτος στόμιου 250-350 mm", 22, PALETTE["accent"], bold=True)
    arrow(d, (620, 535), (980, 535), PALETTE["danger"])
    arrow(d, (980, 535), (620, 535), PALETTE["danger"])
    label(d, (800, 520), "throat 75-85 mm", 21, PALETTE["danger"], bold=True)

    im.save(path)
    return path


def make_collection_hopper():
    path = IMG_DIR / "design-collection-hopper.png"
    im = Image.new("RGB", (1600, 900), PALETTE["white"])
    d = ImageDraw.Draw(im)
    label(d, (80, 60), "Μηχανισμός συλλογής / αποθήκευσης", 34, bold=True, anchor="la")
    label(d, (80, 105), "Hopper με κεκλιμένο πάτο, αναδευτήρα και one-by-one feeder", 24, PALETTE["muted"], anchor="la")

    # hopper shape
    d.polygon([(280, 220), (950, 220), (850, 620), (410, 620)], fill=(245, 248, 250), outline=PALETTE["line"])
    d.line((410, 620, 695, 705, 850, 620), fill=PALETTE["accent"], width=6)
    label(d, (615, 180), "αφαιρούμενο καλάθι / top lid", 24, PALETTE["accent"], bold=True)
    for x, y in [(390, 320), (500, 300), (620, 345), (740, 305), (850, 365), (455, 450), (610, 470), (770, 480), (540, 570), (720, 565)]:
        d.ellipse((x - 42, y - 42, x + 42, y + 42), fill=(210, 235, 80), outline=(130, 150, 30), width=3)

    # agitator
    d.ellipse((550, 420, 690, 560), outline=PALETTE["accent2"], width=6)
    d.line((620, 490, 540, 445), fill=PALETTE["accent2"], width=6)
    d.line((620, 490, 710, 505), fill=PALETTE["accent2"], width=6)
    label(d, (620, 645), "αργός αναδευτήρας anti-jam", 22, PALETTE["accent2"], bold=True)

    # feeder and launcher interface
    rounded(d, (1020, 385, 1210, 575), (255, 255, 255), PALETTE["accent"], 4, 22)
    label(d, (1115, 430), "star feeder", 23, bold=True)
    d.ellipse((1065, 455, 1165, 555), outline=PALETTE["accent"], width=6)
    for ang in [0, 60, 120, 180, 240, 300]:
        import math

        cx, cy = 1115, 505
        ex = cx + 62 * math.cos(math.radians(ang))
        ey = cy + 62 * math.sin(math.radians(ang))
        d.line((cx, cy, ex, ey), fill=PALETTE["accent"], width=5)

    rounded(d, (1280, 390, 1500, 570), (255, 245, 245), PALETTE["danger"], 4, 22)
    label(d, (1390, 430), "προς launcher", 23, PALETTE["danger"], bold=True)
    label(d, (1390, 465), "guide chute", 21, PALETTE["muted"])
    arrow(d, (850, 635), (1020, 505), PALETTE["accent"])
    arrow(d, (1210, 505), (1280, 505), PALETTE["danger"])
    label(d, (840, 735), "service hatch για καθάρισμα μπλοκαρίσματος", 22, PALETTE["muted"])
    im.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "EAF3F5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(20, 92, 102)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=PALETTE["accent"])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    style_table(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*PALETTE["muted"])


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF4DE")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(132, 83, 22)
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(*PALETTE["ink"])
    doc.add_paragraph()


def build_doc():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = [
        make_base_plan(),
        make_side_section(),
        make_pickup_flow(),
        make_intake_front_view(),
        make_collection_hopper(),
    ]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style].font.name = "Arial"
        styles[style].font.color.rgb = RGBColor(*PALETTE["ink"])
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(13)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Tennis Robot - Mechanical Design Concept")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(*PALETTE["accent"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = subtitle.add_run("Βάση, μηχανισμός μαζέματος και μηχανισμός συλλογής/αποθήκευσης")
    rr.font.size = Pt(13)
    rr.font.color.rgb = RGBColor(*PALETTE["muted"])
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Draft concept | 2026-05-18 | Prototype-oriented design")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(*PALETTE["muted"])
    doc.add_picture(str(diagrams[1]), width=Inches(6.9))
    add_caption(doc, "Σχήμα 1. Πλάγια τομή του προτεινόμενου module layout.")
    add_note(
        doc,
        "Κύρια απόφαση",
        "Η εκτόξευση και το μάζεμα παραμένουν χωριστά modules. Η βάση λειτουργεί σαν trolley για μεταφορά, αλλά στην εκτόξευση κάθεται σε stabilizers ώστε να μη στηρίζεται μόνο στους δύο τροχούς.",
    )

    doc.add_heading("1. Design στόχοι", level=1)
    for item in [
        "Καθαρό εξωτερικό σώμα που κρύβει μπαταρία, καλώδια, drive motors, drivers και moving parts.",
        "Μεταφορά τύπου βαλίτσας με δύο μεγάλους τροχούς και τηλεσκοπικό χερούλι.",
        "Σταθερή λειτουργία στην εκτόξευση με μπροστινά stabilizer feet ή φαρδύ rubber foot.",
        "Modular μηχανική: βάση, intake/μάζεμα, hopper/συλλογή και launcher να λύνονται ανεξάρτητα.",
        "Service access για μπλοκαρίσματα, αλλαγή μπαταρίας, καθάρισμα fuzzy debris και πρόσβαση σε electronics.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Design για τη βάση", level=1)
    doc.add_picture(str(diagrams[0]), width=Inches(6.9))
    add_caption(doc, "Σχήμα 2. Κάτοψη βάσης με 2 drive wheels, εσωτερικό σασί και stabilizers.")
    doc.add_paragraph(
        "Η βάση προτείνεται ως εσωτερικό αλουμινένιο σασί με βιδωτά εξωτερικά panels. Τα panels δίνουν καθαρή εμφάνιση και προστασία, αλλά το πραγματικό φορτίο το παίρνει το σασί. Το τηλεσκοπικό χερούλι πρέπει να βιδώνει στο σασί, όχι μόνο στο πλαστικό κέλυφος."
    )
    add_table(
        doc,
        ["Υποσύστημα", "Πρόταση", "Σημείο προσοχής"],
        [
            ["Drive layout", "2 κινητήριοι τροχοί 150-200 mm με DC gear motors/encoders", "Οι τροχοί δίνουν κίνηση και μεταφορά, όχι μόνο στήριξη εκτόξευσης."],
            ["Σταθεροποίηση", "2 πτυσσόμενα stabilizer feet μπροστά ή ένα πλατύ rubber foot", "Να ακουμπά κάτω πριν την εκτόξευση και να σηκώνεται στη μεταφορά."],
            ["Σασί", "2020/2040 extrusion ή aluminum plate 3-5 mm", "Χαμηλή μπαταρία και rigid σημεία για launcher/intake."],
            ["Κέλυφος", "ABS/polycarbonate panels 2-4 mm, βιδωτά", "Service hatches σε intake, electronics και feeder."],
            ["Χερούλι", "Heavy-duty retractable trolley handle", "Αγκύρωση σε δύο κάθετα rails του σασί."],
        ],
        widths=[3.0, 6.3, 6.0],
    )
    add_note(
        doc,
        "Σταθερότητα",
        "Με 2 τροχούς η μεταφορά είναι πρακτική, αλλά για shot repeatability χρειάζεται τρίτο/τέταρτο σημείο επαφής στο έδαφος. Τα stabilizers πρέπει να έχουν λάστιχο για grip και μικρή δυνατότητα ρύθμισης ύψους.",
    )

    doc.add_heading("3. Design για τον μηχανισμό μαζέματος", level=1)
    doc.add_picture(str(diagrams[2]), width=Inches(6.9))
    add_caption(doc, "Σχήμα 3. Μπροστινό intake με funnel, brush roller και elevator.")
    doc.add_picture(str(diagrams[3]), width=Inches(6.9))
    add_caption(doc, "Σχήμα 4. Μπροστινή όψη intake με οδηγούς, στόμιο και brush roller.")
    doc.add_paragraph(
        "Το μάζεμα γίνεται με χαμηλό στόμιο και φαρδείς οδηγούς που συγχωρούν μικρό σφάλμα προσέγγισης. Ένας μαλακός roller/βούρτσα τραβά τη μπάλα προς τα μέσα και τη στέλνει σε μικρή ράμπα ή elevator. Αυτό είναι πιο αξιόπιστο από το να προσπαθήσουμε να χρησιμοποιήσουμε τους flywheels της εκτόξευσης ως intake."
    )
    add_table(
        doc,
        ["Feature", "Starting dimension / spec", "Γιατί"],
        [
            ["Πλάτος στόμιου", "250-350 mm", "Συγχωρεί approach error και bearing noise."],
            ["Ύψος lip", "5-12 mm από το court", "Πιάνει μπάλα χωρίς να ξύνει συνέχεια."],
            ["Brush roller", "60-100 mm διάμετρος, 250-350 mm μήκος", "Μαλακή επαφή που δεν πετάει τη μπάλα μακριά."],
            ["Roller speed", "100-300 rpm αρχικά", "Χρειάζεται tuning με πραγματικές μπάλες."],
            ["Elevator", "Belt/paddle ή inclined roller path", "Ανεβάζει τη μπάλα στο hopper χωρίς χειροκίνητη φόρτωση."],
            ["Sensor", "IR break beam ή microswitch", "Ανιχνεύει ball captured / jam."],
        ],
        widths=[3.6, 5.4, 6.3],
    )

    doc.add_heading("4. Design για τον μηχανισμό συλλογής/αποθήκευσης", level=1)
    doc.add_picture(str(diagrams[4]), width=Inches(6.9))
    add_caption(doc, "Σχήμα 5. Hopper με κεκλιμένο πάτο, anti-jam agitator και star feeder.")
    doc.add_paragraph(
        "Ο μηχανισμός συλλογής είναι το buffer μεταξύ μαζέματος και εκτόξευσης. Θέλει λείες εσωτερικές επιφάνειες, κεκλιμένο πάτο και σημείο καθαρισμού. Για πρώτη πλήρη έκδοση, στόχος 40-80 μπάλες είναι λογικός χωρίς να γίνει υπερβολικά ψηλό το robot."
    )
    add_table(
        doc,
        ["Υποσύστημα", "Πρόταση", "Σημείο προσοχής"],
        [
            ["Hopper", "Κεκλιμένος πάτος προς feeder, λεία ABS/polycarbonate επιφάνεια", "Οι μπάλες να κυλάνε χωρίς γωνίες που μπλοκάρουν."],
            ["Agitator", "Αργός αναδευτήρας ή flexible fingers", "Να σπάει γέφυρες μπαλών χωρίς να τις συμπιέζει."],
            ["Feeder", "Star wheel / paddle wheel για one-by-one έξοδο", "Να μη στέλνει δύο μπάλες στους flywheels."],
            ["Access hatch", "Πλευρικό ή επάνω άνοιγμα", "Καθάρισμα fuzz/debris και ξεμπλοκάρισμα."],
            ["Sensoring", "Ball-present πριν το feeder και jam timeout", "Απαιτείται για ασφαλές launcher control."],
        ],
        widths=[3.2, 6.1, 6.0],
    )

    doc.add_heading("5. Ενοποιημένη διάταξη", level=1)
    add_table(
        doc,
        ["Ζώνη", "Θέση στο robot", "Περιεχόμενο"],
        [
            ["Μπροστά χαμηλά", "Σχεδόν στο court", "Intake lip, funnel guides, brush roller."],
            ["Μπροστά/κέντρο", "Ανεβαίνει προς hopper", "Ράμπα ή elevator, sensor captured."],
            ["Κέντρο επάνω", "Πάνω από μπαταρία", "Hopper 40-80 μπαλών με κεκλιμένο πάτο."],
            ["Πίσω/επάνω", "Προστατευμένο module", "Feeder, dual flywheel launcher, short guarded outlet."],
            ["Κάτω", "Πιο χαμηλά γίνεται", "Μπαταρία, motor drivers, power distribution."],
            ["Πίσω", "Εκτός ball path", "Τηλεσκοπικό χερούλι, charging port, main switch."],
        ],
        widths=[3.5, 4.4, 7.4],
    )

    doc.add_heading("6. Πρώτο BOM επιπέδου design", level=1)
    add_table(
        doc,
        ["Module", "Parts που χρειάζονται για prototype"],
        [
            ["Βάση", "2 drive wheels, 2 DC gear motors με encoders, dual motor driver, 24V battery, aluminum chassis, trolley handle, stabilizer feet, panels, emergency stop, fuse, wiring."],
            ["Μάζεμα", "Funnel plates, brush/foam roller, intake motor, small motor driver, bearings/shaft, adjustable brackets, ramp/elevator belt, ball-present sensor."],
            ["Συλλογή", "Hopper panels, smooth liner, agitator motor, star feeder motor/servo, guide chute, access hatch, jam sensor."],
            ["Εκτόξευση interface", "Reserved space για dual flywheels, guarded outlet, arming switch και RPM sensors. Το αναλυτικό launcher BOM μένει ως ξεχωριστή αγορά."],
        ],
        widths=[3.3, 12.0],
    )

    doc.add_heading("7. Open decisions για CAD", level=1)
    for item in [
        "Τελική χωρητικότητα hopper: 40, 60 ή 80 μπάλες.",
        "Αν το elevator θα είναι belt με paddles ή απλούστερη ράμπα/roller path.",
        "Ακριβής θέση των stabilizers ώστε να μην εμποδίζουν το intake.",
        "Ύψος εξόδου launcher και γωνία εκτόξευσης.",
        "Αν τα panels θα είναι flat-cut ABS για γρήγορο prototype ή 3D printed/fiberglass για πιο τελικό σώμα.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Appendix A - Suggested prototype order", level=1)
    for i, item in enumerate(
        [
            "Φτιάχνουμε mockup βάσης από extrusion/plate και τοποθετούμε τροχούς, μπαταρία, χερούλι.",
            "Προσθέτουμε stabilizer feet και ελέγχουμε αν η βάση μένει ακίνητη σε κραδασμούς.",
            "Φτιάχνουμε intake rig με ρυθμιζόμενο lip/roller gap και δοκιμές με μπάλες στο πάτωμα.",
            "Προσθέτουμε elevator/hopper μικρής χωρητικότητας και sensors.",
            "Κλειδώνουμε τα interface dimensions προς launcher και μετά αγοράζουμε/σχεδιάζουμε flywheel module.",
        ],
        start=1,
    ):
        doc.add_paragraph(f"{i}. {item}")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
